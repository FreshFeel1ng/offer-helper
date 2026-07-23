"""
简历解析与项目知识库

支持 PDF / DOCX / 纯文本简历解析，提取项目经历。
基于提取的内容构建 RAG 知识库，面试时检索相关经历生成针对性回答。
"""
import io
import json
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import numpy as np
from langchain_core.documents import Document

from .config import config

# 简历数据根目录
RESUME_ROOT = Path(__file__).parent.parent / "data" / "resumes"
RESUME_ROOT.mkdir(parents=True, exist_ok=True)


def _get_session_dir(session_id: str) -> Path:
    """获取某个 session 的简历存储目录"""
    session_dir = RESUME_ROOT / session_id
    session_dir.mkdir(parents=True, exist_ok=True)
    return session_dir


@dataclass
class ProjectInfo:
    """项目信息"""
    name: str = ""
    description: str = ""
    role: str = ""
    tech_stack: list[str] = field(default_factory=list)
    highlights: list[str] = field(default_factory=list)
    duration: str = ""
    raw_text: str = ""


@dataclass
class WorkExperience:
    """工作/实习经历"""
    company: str = ""
    role: str = ""
    duration: str = ""
    description: str = ""
    highlights: list[str] = field(default_factory=list)
    tech_stack: list[str] = field(default_factory=list)
    raw_text: str = ""


@dataclass
class ResumeData:
    """简历数据结构"""
    name: str = ""
    summary: str = ""
    skills: list[str] = field(default_factory=list)
    projects: list[ProjectInfo] = field(default_factory=list)
    work_experience: list[WorkExperience] = field(default_factory=list)
    education: list[dict] = field(default_factory=list)
    raw_text: str = ""
    chunks: list[Document] = field(default_factory=list)


class ResumeParser:
    """简历解析器：支持 PDF、DOCX、TXT"""

    def parse(self, content: bytes, filename: str) -> str:
        """根据文件类型解析简历"""
        ext = Path(filename).suffix.lower()
        if ext == ".pdf":
            # 优先用 MinerU 精准解析 API，失败回退本地解析
            if config.mineru_api_token:
                text = self._parse_with_mineru(content, filename)
                if text.strip():
                    return text.strip()
                print("[Resume] MinerU 返回空，回退本地解析")
            return self._parse_pdf(content)
        elif ext in (".docx", ".doc"):
            return self._parse_docx(content)
        else:
            return content.decode("utf-8", errors="ignore")

    def _parse_with_mineru(self, content: bytes, filename: str) -> str:
        """使用 MinerU 精准解析 API"""
        import requests
        import time
        import zipfile

        token = config.mineru_api_token
        base = "https://mineru.net/api/v4"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {token}",
        }

        try:
            print("[Resume] MinerU: 获取上传链接...")
            data = {
                "enable_formula": True,
                "enable_table": True,
                "language": "ch",
                "files": [
                    {
                        "name": filename,
                        "is_ocr": True,
                        "data_id": "resume_0",
                    }
                ],
            }
            resp = requests.post(f"{base}/file-urls/batch", headers=headers, json=data, timeout=30)

            if resp.status_code != 200:
                print(f"[Resume] MinerU 请求失败: {resp.status_code} {resp.text[:200]}")
                return ""

            result = resp.json()
            if result.get("code") != 0:
                print(f"[Resume] MinerU 业务失败: {result.get('msg', '')}")
                return ""

            batch_id = result["data"]["batch_id"]
            file_urls = result["data"]["file_urls"]
            if not file_urls:
                print("[Resume] MinerU: 未获取到上传链接")
                return ""

            print("[Resume] MinerU: 上传文件...")
            upload_resp = requests.put(file_urls[0], data=content, timeout=120)
            if upload_resp.status_code not in (200, 201):
                print(f"[Resume] MinerU 上传失败: {upload_resp.status_code}")
                return ""

            print(f"[Resume] MinerU: batch_id={batch_id}, 等待解析...")

            result_url = f"{base}/extract-results/batch/{batch_id}"
            for attempt in range(40):
                time.sleep(3)
                res = requests.get(result_url, headers=headers, timeout=30)

                if res.status_code != 200 or res.json().get("code") != 0:
                    continue

                data = res.json()["data"]
                extract_results = data.get("extract_result", [])

                if not extract_results:
                    continue

                result_item = extract_results[0]
                state = result_item.get("state", "")

                if state == "done":
                    zip_url = result_item.get("full_zip_url", "")
                    if not zip_url:
                        print("[Resume] MinerU: 无 zip 下载链接")
                        return ""

                    print("[Resume] MinerU: 下载结果...")
                    zip_resp = requests.get(zip_url, timeout=30)
                    if zip_resp.status_code != 200:
                        print(f"[Resume] MinerU 下载失败: {zip_resp.status_code}")
                        return ""

                    with zipfile.ZipFile(io.BytesIO(zip_resp.content)) as zf:
                        for name in zf.namelist():
                            if name.endswith(".md"):
                                text = zf.read(name).decode("utf-8", errors="ignore")
                                print(f"[Resume] MinerU 解析成功: {len(text)} 字符")
                                return text

                    print("[Resume] MinerU: zip 中未找到 md 文件")
                    return ""

                elif state == "failed":
                    print(f"[Resume] MinerU: 任务失败")
                    return ""

                elif attempt % 5 == 0:
                    print(f"[Resume] MinerU: 状态={state}, 继续等待...")

            print("[Resume] MinerU: 超时")
            return ""

        except Exception as e:
            print(f"[Resume] MinerU 异常: {e}")
            return ""

    def _parse_pdf(self, content: bytes) -> str:
        text = ""
        # 方案1: pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"[Resume] pdfplumber 失败: {e}")

        # 方案2: PyPDF2
        if not text.strip():
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(io.BytesIO(content))
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        text += t + "\n"
            except Exception as e:
                print(f"[Resume] PyPDF2 失败: {e}")

        # 方案3: pymupdf (fitz)
        if not text.strip():
            try:
                import fitz
                doc = fitz.open(stream=content, filetype="pdf")
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
            except Exception as e:
                print(f"[Resume] pymupdf 失败: {e}")

        return text.strip()

    def _parse_docx(self, content: bytes) -> str:
        try:
            from docx import Document as DocxDoc
            doc = DocxDoc(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return content.decode("utf-8", errors="ignore")


class ResumeExtractor:
    """从简历文本中提取结构化信息 - 使用 LLM 提取"""

    def extract(self, raw_text: str) -> ResumeData:
        data = ResumeData(raw_text=raw_text)
        data.name = self._extract_name(raw_text)
        data.skills = self._extract_skills(raw_text)
        data.summary = self._build_summary(data)
        # LLM 提取项目经历 和 工作/实习经历
        result = self._extract_all_with_llm(raw_text)
        data.projects = result.get("projects", [])
        data.work_experience = result.get("work_experience", [])
        return data

    def _extract_all_with_llm(self, raw_text: str) -> dict:
        """使用 LLM 从简历中同时提取项目经历和工作/实习经历"""
        from langchain_openai import ChatOpenAI

        text = raw_text[:8000]

        llm = ChatOpenAI(
            model=config.llm_model,
            temperature=0.1,
            api_key=config.deepseek_api_key,
            base_url=config.deepseek_base_url,
        )

        prompt = f"""从以下简历中提取所有项目经历和工作/实习经历。返回严格 JSON 对象。

简历内容：
{text}

返回格式：
{{
  "projects": [
    {{
      "name": "项目名称",
      "description": "项目简要描述（50字内）",
      "role": "你在项目中的角色",
      "tech_stack": ["用到的技术1", "技术2"],
      "highlights": ["亮点1", "亮点2"],
      "duration": "项目时间",
      "raw_text": "简历中关于该项目的原文描述"
    }}
  ],
  "work_experience": [
    {{
      "company": "公司名称",
      "role": "职位",
      "duration": "工作时间段",
      "description": "工作/实习内容总结（100字内）",
      "highlights": ["主要贡献1", "主要贡献2"],
      "tech_stack": ["用到的技术1", "技术2"],
      "raw_text": "简历中关于该段经历的原文描述"
    }}
  ]
}}

注意：
1. 必须返回包含 projects 和 work_experience 两个字段的 JSON 对象
2. 没有的内容返回空数组 []
3. raw_text 必须保留简历中的原文
4. 实习经历也要提取（标记为实习岗位）
5. 不要加任何额外文字或 markdown 代码块标记"""

        try:
            response = llm.invoke(prompt)
            content = response.content if hasattr(response, 'content') else str(response)
            print(f"[Resume] LLM 返回: {content[:300]}...")
            # 清理可能的 markdown 代码块
            content = content.strip()
            if content.startswith("```json"):
                content = content[7:]
            if content.startswith("```"):
                content = content[3:]
            if content.endswith("```"):
                content = content[:-3]
            content = content.strip()

            data = json.loads(content)
            if not isinstance(data, dict):
                print(f"[Resume] LLM 返回格式异常: {type(data)}")
                return {"projects": [], "work_experience": []}

            projects = []
            for p in data.get("projects", []):
                proj = ProjectInfo(
                    name=p.get("name", ""),
                    description=p.get("description", ""),
                    role=p.get("role", ""),
                    tech_stack=p.get("tech_stack", []),
                    highlights=p.get("highlights", []),
                    duration=p.get("duration", ""),
                    raw_text=p.get("raw_text", ""),
                )
                if proj.name or proj.raw_text:
                    projects.append(proj)

            works = []
            for w in data.get("work_experience", []):
                work = WorkExperience(
                    company=w.get("company", ""),
                    role=w.get("role", ""),
                    duration=w.get("duration", ""),
                    description=w.get("description", ""),
                    highlights=w.get("highlights", []),
                    tech_stack=w.get("tech_stack", []),
                    raw_text=w.get("raw_text", ""),
                )
                if work.company or work.raw_text:
                    works.append(work)

            print(f"[Resume] LLM 提取到 {len(projects)} 个项目, {len(works)} 段工作/实习经历")
            return {"projects": projects, "work_experience": works}

        except Exception as e:
            print(f"[Resume] LLM 提取失败，回退到正则: {e}")
            import traceback
            traceback.print_exc()
            projects = self._extract_projects_fallback(raw_text)
            return {"projects": projects, "work_experience": []}

    def _extract_projects_fallback(self, raw_text: str) -> list[ProjectInfo]:
        """正则回退提取"""
        projects = []
        project_sections = self._find_sections(raw_text, [
            "项目经历", "项目经验", "PROJECT", "Projects",
            "项目展示", "个人项目", "主要项目", "工作经验", "工作经历",
        ])
        for section_start, section_end in project_sections:
            section_text = raw_text[section_start:section_end]
            sub_projects = self._split_projects(section_text)
            projects.extend(sub_projects)
        return projects

    def _extract_name(self, text: str) -> str:
        """提取姓名（通常是简历第一行）"""
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        for line in lines[:5]:
            # 中文姓名 2-4 字
            if re.match(r"^[\u4e00-\u9fff]{2,4}$", line):
                return line
            # 英文名
            if re.match(r"^[A-Z][a-z]+(\s[A-Z][a-z]+){0,2}$", line):
                return line
        return ""

    def _extract_skills(self, text: str) -> list[str]:
        """提取技能"""
        skill_keywords = [
            "Python", "Java", "JavaScript", "TypeScript", "Go", "Rust", "C++", "C#",
            "React", "Vue", "Angular", "Node.js", "Django", "Flask", "FastAPI",
            "Spring", "MySQL", "PostgreSQL", "MongoDB", "Redis", "Elasticsearch",
            "Docker", "Kubernetes", "AWS", "Azure", "GCP", "Linux",
            "TensorFlow", "PyTorch", "机器学习", "深度学习", "NLP", "CV",
            "Git", "CI/CD", "Jenkins", "微服务", "分布式", "高并发",
            "HTML", "CSS", "Sass", "Webpack", "Vite", "GraphQL", "REST",
        ]
        found = []
        lower_text = text.lower()
        for skill in skill_keywords:
            if skill.lower() in lower_text:
                found.append(skill)
        return found

    def _find_sections(self, text: str, keywords: list[str]) -> list[tuple[int, int]]:
        """找到简历中的特定段落"""
        sections = []
        next_section_pattern = re.compile(
            r"\n(?:教育|工作|技能|自我|联系方式|基本信息|EDUCATION|WORK|SKILL)",
            re.IGNORECASE
        )

        for kw in keywords:
            pattern = re.compile(rf"(?:^|\n).*?{re.escape(kw)}.*?(?:\n|$)", re.IGNORECASE)
            for match in pattern.finditer(text):
                start = match.start()
                next_match = next_section_pattern.search(text, match.end())
                end = next_match.start() if next_match else min(start + 2000, len(text))
                sections.append((start, end))

        return sections

    def _split_projects(self, section_text: str) -> list[ProjectInfo]:
        """将一个段落拆分为多个项目"""
        projects = []

        parts = re.split(r"\n(?=[\u4e00-\u9fff●■▸▪•·◆◇\-\*\d]+(?:项目|系统|平台|工具|应用))", section_text)

        for part in parts:
            if len(part.strip()) < 20:
                continue

            project = ProjectInfo()
            lines = part.strip().split("\n")

            if lines:
                first_line = re.sub(r"^[\s\-•●■▸]+", "", lines[0]).strip()
                if len(first_line) < 50:
                    project.name = first_line

            project.raw_text = part.strip()

            tech_keywords = [
                "React", "Vue", "Spring", "Django", "Flask", "MySQL", "Redis",
                "Docker", "K8s", "AWS", "Python", "Java", "Go", "Node",
                "微服务", "分布式", "高并发", "机器学习", "深度学习",
            ]
            for tech in tech_keywords:
                if tech.lower() in part.lower():
                    project.tech_stack.append(tech)

            highlights = re.findall(r"[提高降低增加减少优化提升]{1,2}[^\n]*?\d+%?", part)
            project.highlights = highlights[:3]

            role_match = re.search(r"(?:负责|担任|作为|角色)[：:]?\s*(.{2,15}?)(?:[，,\n]|$)", part)
            if role_match:
                project.role = role_match.group(1)

            projects.append(project)

        return projects

    def _build_summary(self, data: ResumeData) -> str:
        """构建简历摘要"""
        parts = []
        if data.name:
            parts.append(f"姓名: {data.name}")
        if data.skills:
            parts.append(f"技能: {', '.join(data.skills[:10])}")
        if data.projects:
            parts.append(f"项目数: {len(data.projects)}个")
        if data.work_experience:
            parts.append(f"工作经历: {len(data.work_experience)}段")
        return " | ".join(parts)


class ResumeKnowledgeBase:
    """简历知识库 - 支持关键词匹配 / 向量 RAG 两种检索模式"""

    MODE_KEYWORD = "keyword"
    MODE_VECTOR = "vector"

    def __init__(self, search_mode: str = MODE_KEYWORD, session_id: str = "default"):
        self.search_mode = search_mode
        self.session_id = session_id
        self.data_dir = _get_session_dir(session_id)
        self.resume: Optional[ResumeData] = None
        self._chunks: list[Document] = []
        self._vector_store = None

    def load_resume(self, content: bytes, filename: str) -> ResumeData:
        """加载并解析简历"""
        parser = ResumeParser()
        extractor = ResumeExtractor()

        print(f"[Resume] 开始解析: {filename}, 大小: {len(content)} bytes")
        raw_text = parser.parse(content, filename)
        print(f"[Resume] 文本提取完成: {len(raw_text)} 字符")

        if not raw_text.strip():
            print("[Resume] 警告: 所有 PDF 解析方式均失败")
            resume = ResumeData()
            self.resume = resume
            return resume

        resume = extractor.extract(raw_text)
        self.resume = resume

        self._build_chunks(resume)
        self._save_local(resume, filename)

        print(f"[Resume] 简历加载完成: 项目={len(resume.projects)}, 技能={len(resume.skills)}")
        return resume

    def load_from_local(self, filename: str) -> Optional[ResumeData]:
        """从本地加载已保存的简历"""
        json_path = self.data_dir / f"{Path(filename).stem}.json"
        if not json_path.exists():
            return None

        with open(json_path, "r", encoding="utf-8") as f:
            data_dict = json.load(f)

        resume = ResumeData(
            name=data_dict.get("name", ""),
            summary=data_dict.get("summary", ""),
            skills=data_dict.get("skills", []),
            raw_text=data_dict.get("raw_text", ""),
        )

        for p in data_dict.get("projects", []):
            resume.projects.append(ProjectInfo(
                name=p.get("name", ""),
                description=p.get("description", ""),
                role=p.get("role", ""),
                tech_stack=p.get("tech_stack", []),
                highlights=p.get("highlights", []),
                duration=p.get("duration", ""),
                raw_text=p.get("raw_text", ""),
            ))

        for w in data_dict.get("work_experience", []):
            resume.work_experience.append(WorkExperience(
                company=w.get("company", ""),
                role=w.get("role", ""),
                duration=w.get("duration", ""),
                description=w.get("description", ""),
                highlights=w.get("highlights", []),
                tech_stack=w.get("tech_stack", []),
                raw_text=w.get("raw_text", ""),
            ))

        self.resume = resume
        self._build_chunks(resume)
        return resume

    def load_latest(self) -> Optional[ResumeData]:
        """自动加载当前 session 最近修改的简历 JSON"""
        json_files = sorted(
            self.data_dir.glob("*.json"),
            key=lambda f: f.stat().st_mtime,
            reverse=True,
        )
        if not json_files:
            legacy_files = sorted(
                RESUME_ROOT.glob("*.json"),
                key=lambda f: f.stat().st_mtime,
                reverse=True,
            )
            if not legacy_files:
                print(f"[Resume] session={self.session_id} 未找到已保存的简历")
                return None
            latest = legacy_files[0]
            print(f"[Resume] 加载旧格式简历: {latest.name}")
        else:
            latest = json_files[0]
            print(f"[Resume] 加载已保存的简历: {latest.name}")
        return self.load_from_local(latest.name)

    def _build_chunks(self, resume: ResumeData):
        """构建知识库块"""
        chunks = []

        if resume.summary:
            chunks.append(Document(
                page_content=f"候选人背景: {resume.summary}",
                metadata={"type": "summary"}
            ))

        if resume.skills:
            chunks.append(Document(
                page_content=f"技能: {', '.join(resume.skills)}",
                metadata={"type": "skills"}
            ))

        for i, proj in enumerate(resume.projects):
            content_parts = [f"项目名称: {proj.name}"]
            if proj.role:
                content_parts.append(f"角色: {proj.role}")
            if proj.tech_stack:
                content_parts.append(f"技术栈: {', '.join(proj.tech_stack)}")
            if proj.highlights:
                content_parts.append(f"亮点: {'; '.join(proj.highlights)}")
            if proj.raw_text:
                content_parts.append(f"详细描述: {proj.raw_text[:500]}")

            chunks.append(Document(
                page_content="\n".join(content_parts),
                metadata={"type": "project", "project_index": i, "project_name": proj.name}
            ))

        for i, exp in enumerate(resume.work_experience):
            content_parts = [f"公司: {exp.company}"]
            if exp.role:
                content_parts.append(f"职位: {exp.role}")
            if exp.duration:
                content_parts.append(f"时间: {exp.duration}")
            if exp.tech_stack:
                content_parts.append(f"技术栈: {', '.join(exp.tech_stack)}")
            if exp.highlights:
                content_parts.append(f"主要贡献: {'; '.join(exp.highlights)}")
            if exp.description:
                content_parts.append(f"工作内容: {exp.description}")
            if exp.raw_text:
                content_parts.append(f"原文: {exp.raw_text[:500]}")

            chunks.append(Document(
                page_content="\n".join(content_parts),
                metadata={"type": "work_experience", "exp_index": i, "company": exp.company}
            ))

        self._chunks = chunks

    def search(self, query: str, top_k: int = 3) -> list[Document]:
        """搜索相关简历内容"""
        if not self._chunks:
            return []

        if self.search_mode == self.MODE_VECTOR:
            return self._search_vector(query, top_k)
        else:
            return self._search_keyword(query, top_k)

    def _search_keyword(self, query: str, top_k: int = 3) -> list[Document]:
        """关键词匹配检索"""
        scored = []
        query_lower = query.lower()
        for chunk in self._chunks:
            content_lower = chunk.page_content.lower()
            score = sum(1 for word in query_lower.split() if word in content_lower)
            if chunk.metadata.get("project_name", "").lower() in query_lower:
                score += 5
            scored.append((score, chunk))

        scored.sort(key=lambda x: x[0], reverse=True)
        return [c for _, c in scored[:top_k] if _ > 0] or [c for _, c in scored[:top_k]]

    def _get_embeddings(self):
        """获取 BGE-M3 嵌入模型"""
        from langchain_openai import OpenAIEmbeddings
        return OpenAIEmbeddings(
            model=config.embedding_model,
            api_key=config.siliconflow_api_key,
            base_url=config.siliconflow_base_url,
        )

    def _init_vector_store(self):
        """初始化 Milvus 向量存储"""
        if self._vector_store is not None:
            return

        from langchain_milvus import Milvus

        embeddings = self._get_embeddings()
        collection_name = f"resume_{self.session_id}_{self.resume.name if self.resume else 'default'}"

        print(f"[RAG] 初始化 Milvus 向量存储: collection={collection_name}")

        self._vector_store = Milvus(
            embedding_function=embeddings,
            collection_name=collection_name,
            connection_args={
                "host": config.milvus_host,
                "port": config.milvus_port,
            },
            auto_id=True,
            drop_old=False,
        )

        if self._chunks:
            texts = [c.page_content for c in self._chunks]
            metadatas = [c.metadata for c in self._chunks]
            print(f"[RAG] 正在向量化 {len(texts)} 个文档块...")
            self._vector_store.add_texts(texts, metadatas=metadatas)
            print(f"[RAG] 向量化完成，已写入 Milvus")

    def _search_vector(self, query: str, top_k: int = 3) -> list[Document]:
        """向量语义检索"""
        try:
            self._init_vector_store()

            if self._vector_store is None:
                print("[RAG] Milvus 未初始化，回退到关键词匹配")
                return self._search_keyword(query, top_k)

            docs = self._vector_store.similarity_search(query, k=top_k)
            print(f"[RAG] 向量检索: query='{query[:30]}...' → {len(docs)} 个结果")
            return docs

        except Exception as e:
            print(f"[RAG] 向量检索失败，回退到关键词匹配: {e}")
            return self._search_keyword(query, top_k)

    def switch_mode(self, mode: str):
        """切换检索模式"""
        if mode not in (self.MODE_KEYWORD, self.MODE_VECTOR):
            raise ValueError(f"不支持的检索模式: {mode}")
        self.search_mode = mode
        print(f"[RAG] 检索模式切换为: {mode}")

        if mode == self.MODE_VECTOR:
            try:
                self._init_vector_store()
            except Exception as e:
                print(f"[RAG] Milvus 初始化失败，保持关键词模式: {e}")
                self.search_mode = self.MODE_KEYWORD

    def get_context_for_question(self, question: str) -> str:
        """根据面试问题检索相关简历上下文 —— 始终返回完整简历"""
        parts = []
        if self.resume:
            if self.resume.summary:
                parts.append(f"候选人背景: {self.resume.summary}")
            if self.resume.skills:
                parts.append(f"技能: {', '.join(self.resume.skills)}")

            for proj in self.resume.projects:
                parts.append(f"项目: {proj.name}\n角色: {proj.role}\n技术栈: {', '.join(proj.tech_stack)}\n亮点: {'; '.join(proj.highlights)}\n描述: {proj.raw_text[:300]}")

            for exp in self.resume.work_experience:
                parts.append(
                    f"工作/实习: {exp.company} | {exp.role} | {exp.duration}\n"
                    f"工作内容: {exp.description}\n"
                    f"技术栈: {', '.join(exp.tech_stack)}\n"
                    f"主要贡献: {'; '.join(exp.highlights)}\n"
                    f"原文: {exp.raw_text[:300]}"
                )
        return "\n\n---\n\n".join(parts) if parts else "暂无候选人简历信息，请根据通用知识回答。"

    def _save_local(self, resume: ResumeData, filename: str):
        """保存简历数据到本地"""
        json_path = self.data_dir / f"{Path(filename).stem}.json"
        data = {
            "name": resume.name,
            "summary": resume.summary,
            "skills": resume.skills,
            "raw_text": resume.raw_text,
            "projects": [
                {
                    "name": p.name,
                    "description": p.description,
                    "role": p.role,
                    "tech_stack": p.tech_stack,
                    "highlights": p.highlights,
                    "duration": p.duration,
                    "raw_text": p.raw_text,
                }
                for p in resume.projects
            ],
            "work_experience": [
                {
                    "company": w.company,
                    "role": w.role,
                    "duration": w.duration,
                    "description": w.description,
                    "highlights": w.highlights,
                    "tech_stack": w.tech_stack,
                    "raw_text": w.raw_text,
                }
                for w in resume.work_experience
            ],
            "education": resume.education,
        }
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
